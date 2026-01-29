from docx import Document
from pathlib import Path
import re

INPUT_DOCX = "Namen gerechten.docx"

OUT_MASTER = "recipes_master.tsv"
OUT_DESCRIPTIONS = "descriptions.txt"
OUT_MAPPING = "mapping.tsv"

doc = Document(INPUT_DOCX)

# 1) haal alle tekst op (ook als Word alles in 1 paragraaf zet)
raw_lines = []
for p in doc.paragraphs:
    if p.text:
        raw_lines.extend(p.text.splitlines())

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text:
                raw_lines.extend(cell.text.splitlines())

rows = []
for line in raw_lines:
    line = line.strip()
    if not line or "|" not in line:
        continue

    left, right = line.split("|", 1)
    kv_id = left.strip().split()[0]
    nl = left.strip()[len(kv_id):].strip()
    en = right.strip()

    if not re.fullmatch(r"KV-\d{3}", kv_id):
        continue

    rows.append((kv_id, nl, en))

# 2) schrijf master.tsv
master_lines = ["id\tnl\ten"] + [f"{kv}\t{nl}\t{en}" for kv, nl, en in rows]
Path(OUT_MASTER).write_text("\n".join(master_lines) + "\n", encoding="utf-8")

# 3) schrijf afgeleiden
Path(OUT_DESCRIPTIONS).write_text("\n".join([en for _, _, en in rows]) + "\n", encoding="utf-8")
Path(OUT_MAPPING).write_text("\n".join([f"{kv}\t{en}" for kv, _, en in rows]) + "\n", encoding="utf-8")

print(f"Klaar: {OUT_MASTER} ({len(rows)} regels)")
print(f"Klaar: {OUT_DESCRIPTIONS} ({len(rows)} regels)")
print(f"Klaar: {OUT_MAPPING} ({len(rows)} regels)")
